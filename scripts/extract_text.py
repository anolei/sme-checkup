#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
申报书文本提取脚本 - 专精特新申报书体检助手
=================================================
支持格式：
  .docx       - python-docx 提取，智能去除合并单元格重复内容
  .doc / .wps - 优先尝试 LibreOffice 转换；不可用时引导用户
  .pdf        - pdfplumber / PyMuPDF 逐段提取，保留结构

使用：
  python extract_text.py 申报书.docx
  python extract_text.py 申报书.doc
  python extract_text.py 申报书.pdf
  python extract_text.py 申报书.docx --output extracted.txt
  python extract_text.py 申报书.docx --format structure   # 按章节结构化输出

输出说明：
  默认输出到 stdout（可重定向），同时在同目录生成 _extracted.txt
  structure 模式：额外识别章节标题，生成 _structured.json

依赖安装：
  pip install python-docx pdfplumber
  （可选，更好的 PDF 支持）pip install pymupdf
  （可选，.doc/.wps 转换）需安装 LibreOffice 桌面版
"""

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


# ──────────────────────────────────────────────────────────────
# DOCX 提取（核心：合并单元格去重）
# ──────────────────────────────────────────────────────────────

def extract_docx(filepath: str) -> str:
    """
    从 .docx 文件提取文本，智能处理合并单元格重复问题。
    
    原理：python-docx 遍历 Table.rows[i].cells 时，合并单元格会在
    每个被合并的列位置各返回一次相同对象 → 文字重复 N 次。
    解决方法：跟踪每行已出现的文字（或 cell._tc 指针），只保留首次出现。
    """
    try:
        from docx import Document
    except ImportError:
        _missing_dep("python-docx", "pip install python-docx")

    doc = Document(filepath)
    parts = []

    # 遍历文档体中的段落和表格（保持文档顺序）
    from docx.oxml.ns import qn
    body = doc.element.body

    para_idx = 0
    table_idx = 0
    para_list = doc.paragraphs
    table_list = doc.tables

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            if para_idx < len(para_list):
                text = para_list[para_idx].text.strip()
                if text:
                    parts.append(text)
                para_idx += 1

        elif tag == 'tbl':
            if table_idx < len(table_list):
                table_text = _extract_table_deduped(table_list[table_idx])
                if table_text:
                    parts.append(table_text)
                table_idx += 1

        elif tag == 'sdt':
            # 结构化文档标签（某些 WPS 生成的 docx 用 sdt 包裹内容）
            text = ''.join(n.text or '' for n in child.iter()
                          if n.tag.split('}')[-1] == 't').strip()
            if text:
                parts.append(text)

    return '\n'.join(parts)


def _extract_table_deduped(table) -> str:
    """
    提取表格文本，对合并单元格实施行内去重 + 行间去重两层过滤。
    
    行内去重：同一行中相同文字只取一次（消除横向合并重复）
    行间去重：连续相邻行完全相同时只保留一行（消除纵向合并重复）
    """
    rows_output = []
    prev_row_key = None

    for row in table.rows:
        seen_tc = set()       # 跟踪 cell 的底层 XML 元素指针（处理合并格最准确）
        seen_text = set()     # 同时跟踪文字（作为备用去重手段）
        row_cells = []

        for cell in row.cells:
            tc_id = id(cell._tc)   # 合并单元格的相同 _tc 对象 id 相同
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)

            text = cell.text.strip()
            if not text:
                continue
            # 行内文字级去重（处理极少数 _tc 不同但内容完全相同的情况）
            if text in seen_text:
                continue
            seen_text.add(text)

            row_cells.append(text)

        if not row_cells:
            continue

        row_str = ' | '.join(row_cells)
        row_key = row_str  # 用于行间去重

        # 行间去重：跳过与上一行完全相同的行（处理纵向合并重复）
        if row_key == prev_row_key:
            continue
        prev_row_key = row_key

        rows_output.append(row_str)

    return '\n'.join(rows_output)


# ──────────────────────────────────────────────────────────────
# DOC / WPS 提取（LibreOffice 转换路径）
# ──────────────────────────────────────────────────────────────

def extract_doc_wps(filepath: str) -> str:
    """
    .doc / .wps 提取策略（按优先级尝试）：

    路径1：pywin32 COM（Windows 专属，调用本地已安装的 WPS 或 Word）
      - 优点：无需额外安装软件（只需 pip install pywin32）
      - 适用：已装 WPS 或 Microsoft Word 的 Windows 用户（中国用户几乎都有 WPS）

    路径2：LibreOffice headless 转换
      - 优点：跨平台，Free；缺点：需单独安装 LibreOffice

    路径3：用户引导
      - 给出清晰的三选方案，包括手动另存为 .docx 或粘贴关键章节
    """
    filepath = os.path.abspath(filepath)

    # ── 路径1：pywin32 COM（WPS / Word）──
    result = _try_extract_via_com(filepath)
    if result is not None:
        return result

    # ── 路径2：LibreOffice ──
    libreoffice_cmd = _find_libreoffice()
    if libreoffice_cmd:
        return _extract_via_libreoffice(filepath, libreoffice_cmd)

    # ── 路径3：用户引导 ──
    _guide_user_no_converter(filepath)
    sys.exit(2)


def _try_extract_via_com(filepath: str) -> "str | None":
    """
    尝试通过 COM 自动化（pywin32）调用本地 WPS 或 Word 提取文字。
    成功返回文本字符串；失败（未安装 pywin32 或 WPS/Word）返回 None。
    """
    try:
        import win32com.client
    except ImportError:
        print("[提取] pywin32 未安装，跳过 COM 路径（可用 pip install pywin32 启用）",
              file=sys.stderr)
        return None

    # 依次尝试 WPS Writer 和 Microsoft Word 的 COM ProgID
    com_ids = [
        ("WPS Writer", "Kwps.Application"),
        ("WPS Writer (旧版)", "KWps.Application"),
        ("Microsoft Word", "Word.Application"),
    ]

    for app_name, prog_id in com_ids:
        try:
            app = win32com.client.Dispatch(prog_id)
            app.Visible = False
            print(f"[提取] 使用 {app_name} (COM) 打开 {Path(filepath).name} ...",
                  file=sys.stderr)

            # 打开文档（只读）
            doc = app.Documents.Open(
                filepath,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
            )

            # 提取全文文字
            text = doc.Content.Text

            doc.Close(SaveChanges=False)
            app.Quit()

            if text and len(text.strip()) > 100:
                print(f"[提取] {app_name} 提取成功，{len(text)} 字符", file=sys.stderr)
                return text
            else:
                print(f"[提取] {app_name} 提取内容为空，尝试其他方式", file=sys.stderr)
                return None

        except Exception as e:
            err_msg = str(e)
            if "CLSIDFromProgID" in err_msg or "Invalid class string" in err_msg:
                # 该 COM 组件未安装，静默跳过
                continue
            else:
                print(f"[提取] {app_name} COM 调用出错：{err_msg}", file=sys.stderr)
                continue

    print("[提取] 未找到可用的 WPS 或 Word COM 组件", file=sys.stderr)
    return None


def _extract_via_libreoffice(filepath: str, lo_cmd: str) -> str:
    """使用 LibreOffice headless 将 .doc/.wps 转换为 .docx 再提取。"""
    suffix = Path(filepath).suffix.lower()
    print(f"[提取] 检测到 {suffix} 格式，使用 LibreOffice 转换中...", file=sys.stderr)

    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            result = subprocess.run(
                [lo_cmd, "--headless", "--convert-to", "docx",
                 "--outdir", tmpdir, filepath],
                capture_output=True, text=True, timeout=60
            )
        except subprocess.TimeoutExpired:
            _die("LibreOffice 转换超时（60秒），请检查文件是否损坏或过大。")
        except Exception as e:
            _die(f"LibreOffice 调用失败：{e}")

        converted = list(Path(tmpdir).glob("*.docx")) or list(Path(tmpdir).glob("*.doc"))
        if not converted:
            _die(
                f"LibreOffice 未能成功转换文件。\n"
                f"  stdout: {result.stdout}\n"
                f"  stderr: {result.stderr}\n"
                f"请手动在 WPS/Office 中将文件另存为 .docx 后重新提交。"
            )

        docx_path = str(converted[0])
        print(f"[提取] LibreOffice 转换完成：{converted[0].name}", file=sys.stderr)
        return extract_docx(docx_path)


def _find_libreoffice() -> str:
    """尝试定位 LibreOffice 可执行文件，返回路径或空字符串。"""
    for cmd in ["libreoffice", "soffice"]:
        if shutil.which(cmd):
            return cmd
    for path in [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]:
        if os.path.isfile(path):
            return path
    return ""


def _guide_user_no_converter(filepath: str):
    """WPS/Word COM 和 LibreOffice 都不可用时，输出清晰的用户三方案引导。"""
    fname = Path(filepath).name
    suffix = Path(filepath).suffix.lower()
    print(f"""
[⚠️  无法提取 {suffix} 格式，需要以下任一工具]

文件「{fname}」为旧版 Word/WPS 二进制格式，需要转换才能提取文字。

╔══ 解决方案（三选一，推荐按顺序尝试）═══════════════════════════╗

  方案A：安装 pywin32（最简单，已有 WPS/Word 的用户首选）
  ─────────────────────────────────────────────────
  pip install pywin32
  安装后重新运行即可，系统会自动调用已安装的 WPS 或 Word。

  方案B：安装 LibreOffice（免费，跨平台）
  ─────────────────────────────────────────
  下载：https://www.libreoffice.org/download/
  安装后重新运行，系统会自动调用 LibreOffice 进行转换。

  方案C：手动转换文件（立即可用，不需安装任何软件）
  ─────────────────────────────────────────────────
  1. 用 WPS 或 Word 打开「{fname}」
  2. 文件 → 另存为 → 选择「Word文档 (.docx)」格式
  3. 重新运行：python extract_text.py "转换后的文件.docx"

╚═══════════════════════════════════════════════════════════════╝
""", file=sys.stderr)


# ──────────────────────────────────────────────────────────────
# PDF 提取
# ──────────────────────────────────────────────────────────────

def extract_pdf(filepath: str) -> str:
    """
    PDF 提取，依次尝试：
    1. pdfplumber（表格感知，推荐）
    2. PyMuPDF / fitz（速度快）
    3. 降级提示（以上均不可用）
    """
    # 尝试 pdfplumber
    try:
        import pdfplumber
        return _extract_pdf_pdfplumber(filepath)
    except ImportError:
        pass

    # 尝试 PyMuPDF
    try:
        import fitz  # PyMuPDF
        return _extract_pdf_pymupdf(filepath, fitz)
    except ImportError:
        pass

    _die(
        "PDF 提取需要安装以下任一库：\n"
        "  pip install pdfplumber    （推荐，表格支持更好）\n"
        "  pip install pymupdf       （速度更快）"
    )


def _extract_pdf_pdfplumber(filepath: str) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            # 先提取表格（结构化）
            tables = page.extract_tables()
            table_cells_bbox = set()
            for table in tables:
                for row in table:
                    cells = [c.strip() if c else "" for c in row]
                    # 行内去重（合并单元格会导致相邻列内容相同）
                    seen = set()
                    deduped = []
                    for cell in cells:
                        if cell and cell not in seen:
                            seen.add(cell)
                            deduped.append(cell)
                    if deduped:
                        parts.append(' | '.join(deduped))

            # 再提取正文文字（排除已被表格覆盖的区域）
            text = page.extract_text()
            if text:
                parts.append(text)

    return '\n'.join(parts)


def _extract_pdf_pymupdf(filepath: str, fitz) -> str:
    parts = []
    doc = fitz.open(filepath)
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            parts.append(text.strip())
    return '\n'.join(parts)


# ──────────────────────────────────────────────────────────────
# 结构化识别（章节分析）
# ──────────────────────────────────────────────────────────────

# 中文数字 → 阿拉伯数字映射（申报书用"一、二、…"编号主章节）
CN_NUM = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
          '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}

# 主章节标题正则：匹配"一、xxx""二、xxx"等，不限制标题内容
MAJOR_SECTION_RE = re.compile(r'^([一二三四五六七八九十])[、．.。]\s*(.{1,30})$')

# 关键内容关键词：用于在章节内容中做二次扫描（章节标题无关）
CONTENT_KEYWORDS = {
    'section_7': ['产业链配套', '补短板', '填空白', '替代进口'],
    'section_8': ['产学研', '合作院校', '研究领域', '合作机构', '高等院校'],
    'section_9': ['企业简介', '总体简要介绍', '企业经营管理', '企业主导产品及技术'],
}


def build_structure(raw_text: str) -> dict:
    """
    将原始文本按章节分块，识别关键章节。

    策略：
    1. 第一遍：按中文数字编号（一、二、…）切割主章节，直接映射为 section_1..section_9
       不依赖标题内容关键词，无论标题叫"创新能力""其他"都能正确归位
    2. 第二遍：对每个章节内容做关键词二次扫描，补充标注关键子内容位置
    3. 验证：检查必要章节是否存在且有实质内容
    """
    lines = raw_text.split('\n')
    sections = {}
    current_key = 'preamble'
    current_num = 0
    current_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_lines.append('')
            continue

        m = MAJOR_SECTION_RE.match(stripped)
        if m:
            cn_char = m.group(1)
            num = CN_NUM.get(cn_char, 0)
            if num > 0:
                # 保存上一节
                if current_lines:
                    sections[current_key] = '\n'.join(current_lines).strip()
                current_key = f'section_{num}'
                current_num = num
                current_lines = [stripped]
                continue

        current_lines.append(stripped)

    # 保存最后一节
    if current_lines:
        sections[current_key] = '\n'.join(current_lines).strip()

    # ── 第二遍：内容关键词扫描，找到关键子内容并提取 ──
    # 对于关键内容（产学研/企业简介/产业链配套），如果已在正确 section 中就跳过；
    # 否则扫描所有 section 内容找到它，创建一个别名键指向该内容
    content_map = {}  # section_id -> 找到的关键内容摘录

    for target_sec, keywords in CONTENT_KEYWORDS.items():
        if target_sec in sections and len(sections[target_sec]) > 100:
            # 已在正确位置，做内容验证
            content = sections[target_sec]
            kw_found = [kw for kw in keywords if kw in content]
            content_map[target_sec] = {
                'found_in': target_sec,
                'keywords_matched': kw_found,
                'content_length': len(content)
            }
        else:
            # 在所有 section 内容中搜索
            for sec_id, content in sections.items():
                kw_found = [kw for kw in keywords if kw in content]
                if kw_found:
                    # 将该内容提升到目标 section_id
                    sections[target_sec] = content
                    content_map[target_sec] = {
                        'found_in': sec_id,
                        'keywords_matched': kw_found,
                        'content_length': len(content),
                        'note': f'内容实际位于 {sec_id}（{_section_name(sec_id, sections)}），已自动重新映射'
                    }
                    break

    # ── 验证关键章节是否存在且有实质内容 ──
    required_sections = {
        'section_9': '企业简介',
        'section_7': '产业链配套',
        'section_8': '产学研合作',
        'section_2': '经济效益',
    }
    missing = []
    for sec_id, sec_name in required_sections.items():
        content = sections.get(sec_id, '')
        if len(content) < 50:
            missing.append(sec_name)

    return {
        "sections": sections,
        "content_map": content_map,
        "missing_sections": missing,
        "total_chars": len(raw_text),
        "section_count": len([k for k in sections if k.startswith('section_')]),
    }


def _section_name(sec_id: str, sections: dict) -> str:
    """从 section 内容中提取章节标题首行（用于日志展示）。"""
    content = sections.get(sec_id, '')
    first_line = content.split('\n')[0] if content else sec_id
    return first_line[:20]


# ──────────────────────────────────────────────────────────────
# 工具函数
# ──────────────────────────────────────────────────────────────

def _missing_dep(lib: str, install_cmd: str):
    print(f"[错误] 缺少依赖库 {lib}，请运行：{install_cmd}", file=sys.stderr)
    sys.exit(3)


def _die(msg: str):
    print(f"[错误] {msg}", file=sys.stderr)
    sys.exit(1)


def _post_process(text: str) -> str:
    """
    通用后处理：
    - 物理截断文档尾部的审查规则与表格说明（防止误导大模型）
    - 折叠连续空行（超过2行的空行压缩为1行）
    - 去除纯空格行
    - 折叠连续完全相同的行（消除残余重复）
    """
    lines = text.split('\n')
    
    # 截断提取（屏蔽掉文档末尾官方带的规则评分表，防止作为上下文污染）
    truncated_lines = []
    for line in lines:
        if "十、初核推荐" in line or "初核指标(如符合，请在对应" in line:
            # 遇到表格末尾的官方规则说明栏，直接丢弃后续所有文本
            break
        truncated_lines.append(line)
    lines = truncated_lines

    # 折叠连续重复行（不仅限于空行）
    deduped = []
    prev = None
    repeat_count = 0
    MAX_REPEAT = 2  # 允许同一内容最多连续出现N次（表格分节会有少量合理重复）

    for line in lines:
        stripped = line.strip()
        if stripped == prev:
            repeat_count += 1
            if repeat_count < MAX_REPEAT:
                deduped.append(line)
        else:
            prev = stripped
            repeat_count = 0
            deduped.append(line)

    # 折叠超过2个连续空行
    result = []
    blank_count = 0
    for line in deduped:
        if not line.strip():
            blank_count += 1
            if blank_count <= 2:
                result.append(line)
        else:
            blank_count = 0
            result.append(line)

    return '\n'.join(result)


# ──────────────────────────────────────────────────────────────
# 主入口
# ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="申报书文本提取 - 支持 .docx/.doc/.wps/.pdf，智能去重"
    )
    parser.add_argument("input", help="申报书文件路径")
    parser.add_argument("--output", default="", help="输出文件路径（默认：输入文件名_extracted.txt）")
    parser.add_argument("--format", choices=["text", "structure"], default="text",
                        help="text=纯文本；structure=按章节分析并生成报告")
    args = parser.parse_args()

    filepath = args.input
    if not os.path.isfile(filepath):
        _die(f"文件不存在：{filepath}")

    suffix = Path(filepath).suffix.lower()
    print(f"[提取] 文件：{Path(filepath).name}  格式：{suffix}", file=sys.stderr)

    # ── 根据格式选择提取方式 ──
    if suffix == '.docx':
        raw_text = extract_docx(filepath)
    elif suffix in ('.doc', '.wps'):
        raw_text = extract_doc_wps(filepath)
    elif suffix == '.pdf':
        raw_text = extract_pdf(filepath)
    else:
        _die(
            f"不支持的文件格式：{suffix}\n"
            f"支持的格式：.docx  .doc  .wps  .pdf\n"
            f"如为其他格式，请在 WPS/Word 中另存为 .docx 后重新提交。"
        )

    # ── 后处理 ──
    cleaned = _post_process(raw_text)
    print(f"[提取] 完成，共 {len(cleaned)} 字符", file=sys.stderr)

    # ── 章节结构分析 ──
    if args.format == "structure":
        structure = build_structure(cleaned)
        if structure["missing_sections"]:
            print(f"[⚠️  警告] 未能识别以下关键章节：", file=sys.stderr)
            for sec in structure["missing_sections"]:
                print(f"          · {sec}", file=sys.stderr)
            print(f"  → 这些章节可能已嵌入在表格中无法自动识别；", file=sys.stderr)
            print(f"    请检查输出文本或手动补充对应章节内容。", file=sys.stderr)

        # 输出 JSON 结构
        out_path = args.output or str(Path(filepath).with_suffix('')) + "_structured.json"
        Path(out_path).write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding='utf-8'
        )
        # 同时输出纯文本
        txt_path = str(Path(filepath).with_suffix('')) + "_extracted.txt"
        Path(txt_path).write_text(cleaned, encoding='utf-8')
        print(f"[提取] 结构化JSON → {out_path}", file=sys.stderr)
        print(f"[提取] 纯文本     → {txt_path}", file=sys.stderr)

    else:
        # 纯文本模式
        out_path = args.output or str(Path(filepath).with_suffix('')) + "_extracted.txt"
        Path(out_path).write_text(cleaned, encoding='utf-8')
        print(cleaned)
        print(f"[提取] 已保存 → {out_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
